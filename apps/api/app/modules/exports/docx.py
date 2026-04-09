from io import BytesIO
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[attr-defined]

# Brand colour: #1D9E75
BRAND_RGB = RGBColor(0x1D, 0x9E, 0x75)

SECTION_LABELS: dict[str, str] = {
    "resumenEjecutivo": "Resumen Ejecutivo",
    "problema": "El Problema",
    "solucion": "Nuestra Solución",
    "alcance": "Alcance del Proyecto",
    "timeline": "Cronograma",
    "inversion": "Inversión",
    "proximosPasos": "Próximos Pasos",
}
SECTION_ORDER = [
    "resumenEjecutivo",
    "problema",
    "solucion",
    "alcance",
    "timeline",
    "inversion",
    "proximosPasos",
]


def _apply_brand_color(run: object) -> None:
    """Set run font colour to the brand green."""
    run.font.color.rgb = BRAND_RGB  # type: ignore[attr-defined]


def generate_docx(
    sections: dict[str, str],
    client_name: str,
    company: str,
    title: str = "",
) -> bytes:
    """
    Genera DOCX usando python-docx.

    Layout:
      - Cover page: proposal title, client name / company, date
      - Each section as Heading 2 (brand colour) + body paragraph
    """
    doc = Document()

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------
    cover_title = title or f"Propuesta Comercial — {company}"

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(cover_title)
    title_run.bold = True
    title_run.font.size = Pt(26)
    _apply_brand_color(title_run)

    doc.add_paragraph()  # spacer

    if client_name or company:
        recipient_label = company if company else client_name
        if client_name and company:
            recipient_label = f"{client_name} — {company}"
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(f"Preparada para: {recipient_label}")
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(14)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(date.today().strftime("%d de %B de %Y"))
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ------------------------------------------------------------------
    # Sections
    # ------------------------------------------------------------------
    for key in SECTION_ORDER:
        content = sections.get(key, "")
        if not content:
            continue

        heading_para = doc.add_paragraph()
        heading_run = heading_para.add_run(SECTION_LABELS[key])
        heading_run.bold = True
        heading_run.font.size = Pt(14)
        _apply_brand_color(heading_run)

        doc.add_paragraph(content)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
